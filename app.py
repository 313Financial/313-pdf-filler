import base64
import io
import os
import requests
from flask import Flask, request, jsonify
from pypdf import PdfReader, PdfWriter
from docx import Document
from docx.shared import Pt
import copy
from datetime import date

app = Flask(__name__)

CONCIERGE_TEMPLATE_URL = "https://raw.githubusercontent.com/313Financial/313-pdf-filler/main/Concierge%20Form%20May%2025%20v2%20editable.pdf"
DIP_TEMPLATE_URL = "https://raw.githubusercontent.com/313Financial/313-pdf-filler/main/DIP_CERT_TEMPLATE.docx"

FIELD_MAP = {
    "borrower_name":       "Borrower name",
    "rate_product":        "Rate / product",
    "use_of_funds":        "Use of funds",
    "charge_type":         "charge type",
    "security_address":    "Sales particulars",
    "serviced_or_retained":"Serviced or retained",
    "dual_rep":            "Dual rep",
    "sols_email":          "Sols email",
    "exit_strategy":       "Exit strategy",
    "gross_loan":          "gross loan",
    "net_loan":            "net loan",
    "loan_term":           "Term",
    "broker_fee":          "Broker fee",
}

CHECKBOX_MAP = {
    "rate_product":        "Rate",
    "use_of_funds":        "Use",
    "charge_type":         "Charge",
    "serviced_or_retained":"Serviced",
    "sols_email":          "Sol email",
}

@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})


@app.route("/fill-pdf", methods=["POST"])
def fill_pdf():
    try:
        data = request.get_json(force=True)
        if not data:
            return jsonify({"error": "No JSON body"}), 400

        response = requests.get(CONCIERGE_TEMPLATE_URL, timeout=30)
        if response.status_code != 200:
            return jsonify({"error": f"Failed to download template: {response.status_code}"}), 500

        template_buf = io.BytesIO(response.content)
        reader = PdfReader(template_buf)
        writer = PdfWriter()
        writer.append(reader)

        fields = {}
        for json_key, pdf_field in FIELD_MAP.items():
            value = data.get(json_key, "")
            if value:
                fields[pdf_field] = str(value)

        for json_key, checkbox_field in CHECKBOX_MAP.items():
            if data.get(json_key, ""):
                fields[checkbox_field] = "/Yes"

        writer.update_page_form_field_values(writer.pages[0], fields, auto_regenerate=False)

        buf = io.BytesIO()
        writer.write(buf)
        pdf_bytes = buf.getvalue()
        pdf_base64 = base64.b64encode(pdf_bytes).decode("utf-8")

        borrower = data.get("borrower_name", "Unknown").replace(" ", "_")
        filename = f"Together_Concierge_{borrower}.pdf"

        return jsonify({"pdf_base64": pdf_base64, "filename": filename})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/generate-dip", methods=["POST"])
def generate_dip():
    try:
        data = request.get_json(force=True)
        if not data:
            return jsonify({"error": "No JSON body"}), 400

        customer_names = data.get("customer_names", "")
        company_name = data.get("company_name", "N/A")
        loan_amount = data.get("loan_amount", "")
        adviser_name = data.get("adviser_name", "Paul Gray")
        adviser_email = data.get("adviser_email", "paul@313group.co.uk")
        adviser_phone = data.get("adviser_phone", "01912286969")
        decision_date = date.today().strftime("%d/%m/%Y")

        response = requests.get(DIP_TEMPLATE_URL, timeout=30)
        if response.status_code != 200:
            return jsonify({"error": f"Failed to download DIP template: {response.status_code}"}), 500

        doc = Document(io.BytesIO(response.content))

        replacements = {
            "{{customer_names}}": customer_names,
            "{{company_name}}": company_name,
            "{{loan_amount}}": loan_amount,
            "{{adviser_name}}": adviser_name,
            "{{adviser_email}}": adviser_email,
            "{{adviser_phone}}": adviser_phone,
            "{{decision_date}}": decision_date,
        }

        def replace_in_paragraph(para):
            for key, value in replacements.items():
                if key in para.text:
                    for run in para.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)

        for para in doc.paragraphs:
            replace_in_paragraph(para)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para)

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
        docx_base64 = base64.b64encode(docx_bytes).decode("utf-8")

        safe_name = customer_names.replace(" ", "_")
        filename = f"DIP_Certificate_{safe_name}.docx"

        return jsonify({"pdf_base64": docx_base64, "filename": filename})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
